import os
import sys
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

# Colores Ejecutivos
COLOR_PRIMARY = RGBColor(26, 36, 56)      # Navy #1A2438
COLOR_SECONDARY = RGBColor(180, 115, 30)  # Oro/Bronce #B4731E
COLOR_DARK = RGBColor(45, 55, 72)         # Gris oscuro #2D3748
COLOR_MUTED = RGBColor(113, 128, 150)     # Gris medio #718096
COLOR_ACCENT = RGBColor(197, 48, 48)      # Rojo Alerta #C53030

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=140, right=140):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_formatted_runs(paragraph, text, default_color=COLOR_DARK, default_size=Pt(10), is_bullet=False):
    """Parsea markdown inline (**negrita**, *cursiva*, `código`) y añade runs limpios."""
    # Eliminar posibles marcas de listas o comillas de citación residuales
    clean_text = text
    if is_bullet:
        clean_text = re.sub(r'^[\*\-\+]\s+', '', clean_text)
    
    # Tokenizar para negrita y cursiva
    # Regex para capturar **negrita**, *cursiva*, `código`
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    tokens = re.split(pattern, clean_text)
    
    for token in tokens:
        if not token:
            continue
        
        bold = False
        italic = False
        font_name = "Arial"
        font_size = default_size
        color = default_color
        t_text = token

        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            bold = True
            t_text = token[2:-2]
            color = COLOR_PRIMARY
        elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
            italic = True
            t_text = token[1:-1]
            color = COLOR_SECONDARY
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            t_text = token[1:-1]
            font_name = "Consolas"
            font_size = Pt(9)
            color = COLOR_DARK

        # Limpieza de corchetes de enlaces tipo [[Nota]] o [Texto](url)
        t_text = re.sub(r'\[\[(.*?)\]\]', r'\1', t_text)
        t_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', t_text)

        r = paragraph.add_run(t_text)
        r.font.name = font_name
        r.font.size = font_size
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color

def parse_markdown_file_to_docx(doc, filepath):
    """Parsea completamente un archivo Markdown e inyecta elementos limpios en el documento Word."""
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()

    lines = content.split('\n')
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal table_rows, in_table
        if table_rows:
            # Filtrar filas vacías o malformadas
            valid_rows = [r for r in table_rows if len(r) > 1 or (len(r) == 1 and r[0] != '')]
            if valid_rows:
                col_count = max(len(r) for r in valid_rows)
                t = doc.add_table(rows=len(valid_rows), cols=col_count)
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                t.autofit = True
                
                for r_i, row in enumerate(valid_rows):
                    for c_i in range(col_count):
                        val = row[c_i] if c_i < len(row) else ""
                        cell = t.cell(r_i, c_i)
                        set_cell_margins(cell, top=100, bottom=100, left=130, right=130)
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        
                        if r_i == 0:
                            set_cell_background(cell, "1A2438")
                            add_formatted_runs(p, val, default_color=RGBColor(255, 255, 255), default_size=Pt(9))
                            for r in p.runs:
                                r.font.bold = True
                        else:
                            if r_i % 2 == 1:
                                set_cell_background(cell, "F8F9FA")
                            add_formatted_runs(p, val, default_color=COLOR_DARK, default_size=Pt(8.5))
                doc.add_paragraph()
            table_rows = []
            in_table = False

    for line in lines:
        raw = line.strip()

        # 1. Manejo de bloques de código / cajas ASCII ```
        if raw.startswith('```'):
            if in_code_block:
                # Cerrar bloque de código: renderizar en una caja sombreada
                if code_block_lines:
                    t_box = doc.add_table(rows=1, cols=1)
                    t_box.alignment = WD_TABLE_ALIGNMENT.CENTER
                    cell = t_box.cell(0, 0)
                    set_cell_background(cell, "F1F5F9") # Slate claro
                    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
                    p_box = cell.paragraphs[0]
                    p_box.paragraph_format.space_before = Pt(0)
                    p_box.paragraph_format.space_after = Pt(0)
                    
                    full_code = '\n'.join(code_block_lines)
                    r_code = p_box.add_run(full_code)
                    r_code.font.name = "Consolas"
                    r_code.font.size = Pt(8.5)
                    r_code.font.color.rgb = COLOR_DARK
                    doc.add_paragraph()
                code_block_lines = []
                in_code_block = False
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(line)
            continue

        # 2. Manejo de Tablas Markdown
        if raw.startswith('|') and raw.endswith('|'):
            # Ignorar separadores |---|---|
            if re.match(r'^\|[\s\-:|]+\|$', raw):
                continue
            cells = [c.strip() for c in raw.split('|')[1:-1]]
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            flush_table()

        if not raw:
            continue

        # 3. Encabezados H1, H2, H3
        if raw.startswith('# '):
            h_text = raw.replace('# ', '').strip()
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(22)
            h.paragraph_format.space_after = Pt(6)
            add_formatted_runs(h, h_text, default_color=COLOR_PRIMARY, default_size=Pt(16))
            for r in h.runs:
                r.font.bold = True
        elif raw.startswith('## '):
            h_text = raw.replace('## ', '').strip()
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(4)
            add_formatted_runs(h, h_text, default_color=COLOR_SECONDARY, default_size=Pt(13))
            for r in h.runs:
                r.font.bold = True
        elif raw.startswith('### '):
            h_text = raw.replace('### ', '').strip()
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(3)
            add_formatted_runs(h, h_text, default_color=COLOR_PRIMARY, default_size=Pt(11))
            for r in h.runs:
                r.font.bold = True

        # 4. Citas / Callouts (>)
        elif raw.startswith('> '):
            q_text = raw.replace('> ', '').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, q_text, default_color=COLOR_SECONDARY, default_size=Pt(9.5))
            for r in p.runs:
                r.font.italic = True

        # 5. Listas con viñetas (- o *)
        elif raw.startswith('- ') or raw.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_runs(p, raw, default_color=COLOR_DARK, default_size=Pt(9.5), is_bullet=True)

        # 6. Listas numeradas (1. 2. etc)
        elif re.match(r'^\d+\.\s+', raw):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            item_text = re.sub(r'^\d+\.\s+', '', raw)
            add_formatted_runs(p, item_text, default_color=COLOR_DARK, default_size=Pt(9.5))

        # 7. Separadores horizontales (---)
        elif raw in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run("─" * 45)
            r.font.color.rgb = COLOR_MUTED
            r.font.size = Pt(8)

        # 8. Párrafos normales
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            add_formatted_runs(p, raw, default_color=COLOR_DARK, default_size=Pt(9.5))

    if in_table:
        flush_table()

def generate_perfect_alby_docx():
    doc = Document()
    
    # Márgenes de 1 pulgada
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 1. PORTADA EDITORIAL IMPECABLE
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(30)

    p_badge = doc.add_paragraph()
    r_badge = p_badge.add_run("💎 ARQUEÓLOGO CREATIVO / CREATIVE ARCHAEOLOGIST")
    r_badge.font.name = "Arial"
    r_badge.font.size = Pt(11)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_SECONDARY

    p_tit = doc.add_paragraph()
    r_tit = p_tit.add_run("INFORME MAESTRO DE AUDITORÍA FORENSE & RESCATE DE ACTIVOS")
    r_tit.font.name = "Arial"
    r_tit.font.size = Pt(22)
    r_tit.font.bold = True
    r_tit.font.color.rgb = COLOR_PRIMARY
    p_tit.paragraph_format.space_after = Pt(8)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Dictamen Estratégico, Desentierro de IPs, Radiografía Financiera y Roadmap 2026–2027")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_MUTED
    p_sub.paragraph_format.space_after = Pt(35)

    # Tabla de Metadatos de Portada
    t_meta = doc.add_table(rows=4, cols=2)
    t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_info = [
        ("AUTOR AUDITADO:", "Alby Ojeda (Guionista, Diseñador Narrativo & Creador)"),
        ("TRIBUNAL AUDITOR:", "Panel de 6 Agentes (IP Scout, Crítico, Buyer Persona, Hater, Pricing, Psicólogo)"),
        ("FECHA DE EMISIÓN:", "21 de Agosto de 2026"),
        ("ALCANCE DEL ANÁLISIS:", "+300.000 archivos | 522 guiones | 10 años de facturas históricas")
    ]
    for idx, (lbl, val) in enumerate(meta_info):
        c1, c2 = t_meta.cell(idx, 0), t_meta.cell(idx, 1)
        set_cell_margins(c1, 80, 80, 120, 120)
        set_cell_margins(c2, 80, 80, 120, 120)
        set_cell_background(c1, "EDF2F7")
        set_cell_background(c2, "F7FAFC")
        c1.paragraphs[0].paragraph_format.space_after = Pt(0)
        c2.paragraphs[0].paragraph_format.space_after = Pt(0)
        
        r1 = c1.paragraphs[0].add_run(lbl)
        r1.font.name = "Arial"
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = COLOR_PRIMARY
        
        r2 = c2.paragraphs[0].add_run(val)
        r2.font.name = "Arial"
        r2.font.size = Pt(9)
        r2.font.color.rgb = COLOR_DARK

    doc.add_page_break()

    # 2. INGESTA Y PARSEO LIMPIO DE LOS 5 INFORMES MAESTROS
    sources = [
        r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\Estrategia\00_INFORME_MAESTRO_DEFINITIVO_Y_CALIBRADO.md",
        r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\Estrategia\02_LISTA_DE_PRECIOS_REALISTA_2026.md",
        r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\Estrategia\04_GUIA_MAESTRA_FILOSOFIA_LABORAL_Y_FILTRADO.md",
        r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\05_DICTAMEN_DEMOLEDOR_CONSEJO_MAESTRO.md",
        r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\2 Nivel Tactica\06_ROADMAP_ESTRATEGICO_PRODUCTOS_2026_2027.md"
    ]

    for s in sources:
        if os.path.exists(s):
            parse_markdown_file_to_docx(doc, s)
            doc.add_page_break()

    output_path = r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\INFORME_MAESTRO_ARQUEOLOGO_CREATIVO_ALBY_OJEDA.docx"
    doc.save(output_path)
    print(f"✅ DOCUMENTO WORD PERFECCIONADO (SIN MARKDOWN CRUDO): {output_path}")

    # También actualizar el script en el repositorio
    repo_script = r"D:\PROYECTOS\VibeCoding\Arqueologo-Creativo\scripts\generate_docx_report.py"
    with open(repo_script, 'w', encoding='utf-8') as fp:
        with open(__file__, 'r', encoding='utf-8') as this_fp:
            fp.write(this_fp.read())
    print(f"✅ Script del repositorio actualizado con el motor de parseo avanzado.")

if __name__ == "__main__":
    generate_perfect_alby_docx()
