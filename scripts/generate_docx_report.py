import os
import sys
import argparse
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

COLOR_PRIMARY = RGBColor(26, 36, 56)      # Navy #1A2438
COLOR_SECONDARY = RGBColor(180, 115, 30)  # Gold/Bronze #B4731E
COLOR_DARK = RGBColor(45, 55, 72)         # Dark Gray #2D3748
COLOR_MUTED = RGBColor(113, 128, 150)     # Medium Gray #718096

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=140, right=140):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def clean_latex_formula(text):
    """Converts raw LaTeX math expressions into clean readable text."""
    t = text
    t = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1) / (\2)', t)
    t = re.sub(r'\\text\{([^}]+)\}', r'\1', t)
    t = t.replace(r'\times', '×').replace(r'\le', '≤').replace(r'\ge', '≥')
    t = t.replace('$$', '').replace('$', '').strip()
    return t

def add_formatted_runs(paragraph, text, default_color=COLOR_DARK, default_size=Pt(10), is_bullet=False):
    clean_text = text
    if is_bullet:
        clean_text = re.sub(r'^[\*\-\+]\s+', '', clean_text)
    
    if '$' in clean_text:
        clean_text = clean_latex_formula(clean_text)

    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    tokens = re.split(pattern, clean_text)
    
    for token in tokens:
        if not token:
            continue
        
        bold = False
        italic = False
        font_name = "Arial"
        font_size = default_size
        color = default_color
        t_text = token

        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            bold = True
            t_text = token[2:-2]
            color = COLOR_PRIMARY
        elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
            italic = True
            t_text = token[1:-1]
            color = COLOR_SECONDARY
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            t_text = token[1:-1]
            font_name = "Consolas"
            font_size = Pt(9)
            color = COLOR_DARK

        t_text = re.sub(r'\[\[(.*?)\]\]', r'\1', t_text)
        t_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', t_text)

        r = paragraph.add_run(t_text)
        r.font.name = font_name
        r.font.size = font_size
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color

def parse_markdown_to_docx(doc, filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()

    lines = content.split('\n')
    in_code_block = False
    is_mermaid_block = False
    code_block_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal table_rows, in_table
        if table_rows:
            valid_rows = [r for r in table_rows if len(r) > 1 or (len(r) == 1 and r[0] != '')]
            if valid_rows:
                col_count = max(len(r) for r in valid_rows)
                t = doc.add_table(rows=len(valid_rows), cols=col_count)
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                t.autofit = True
                
                for r_i, row in enumerate(valid_rows):
                    for c_i in range(col_count):
                        val = row[c_i] if c_i < len(row) else ""
                        cell = t.cell(r_i, c_i)
                        set_cell_margins(cell, top=100, bottom=100, left=130, right=130)
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        
                        if r_i == 0:
                            set_cell_background(cell, "1A2438")
                            add_formatted_runs(p, val, default_color=RGBColor(255, 255, 255), default_size=Pt(9))
                            for r in p.runs:
                                r.font.bold = True
                        else:
                            if r_i % 2 == 1:
                                set_cell_background(cell, "F8F9FA")
                            add_formatted_runs(p, val, default_color=COLOR_DARK, default_size=Pt(8.5))
                doc.add_paragraph()
            table_rows = []
            in_table = False

    for line in lines:
        raw = line.strip()

        if raw.startswith('```'):
            if in_code_block:
                if code_block_lines and not is_mermaid_block:
                    t_box = doc.add_table(rows=1, cols=1)
                    t_box.alignment = WD_TABLE_ALIGNMENT.CENTER
                    cell = t_box.cell(0, 0)
                    set_cell_background(cell, "F1F5F9")
                    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
                    p_box = cell.paragraphs[0]
                    p_box.paragraph_format.space_before = Pt(0)
                    p_box.paragraph_format.space_after = Pt(0)
                    
                    full_code = '\n'.join(code_block_lines)
                    r_code = p_box.add_run(full_code)
                    r_code.font.name = "Consolas"
                    r_code.font.size = Pt(8.5)
                    r_code.font.color.rgb = COLOR_DARK
                    doc.add_paragraph()
                code_block_lines = []
                in_code_block = False
                is_mermaid_block = False
            else:
                if in_table:
                    flush_table()
                in_code_block = True
                if 'mermaid' in raw.lower():
                    is_mermaid_block = True
            continue

        if in_code_block:
            if not is_mermaid_block:
                code_block_lines.append(line)
            continue

        if raw.startswith('$$') and raw.endswith('$$'):
            clean_form = clean_latex_formula(raw)
            p_form = doc.add_paragraph()
            p_form.paragraph_format.left_indent = Inches(0.5)
            p_form.paragraph_format.space_before = Pt(6)
            p_form.paragraph_format.space_after = Pt(6)
            r_form = p_form.add_run("📐 " + clean_form)
            r_form.font.name = "Arial"
            r_form.font.size = Pt(10)
            r_form.font.bold = True
            r_form.font.color.rgb = COLOR_SECONDARY
            continue

        if raw.startswith('|') and raw.endswith('|'):
            if re.match(r'^\|[\s\-:|]+\|$', raw):
                continue
            cells = [c.strip() for c in raw.split('|')[1:-1]]
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            flush_table()

        if not raw:
            continue

        if raw.startswith('# '):
            h_text = raw.replace('# ', '').strip()
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(22)
            h.paragraph_format.space_after = Pt(6)
            add_formatted_runs(h, h_text, default_color=COLOR_PRIMARY, default_size=Pt(16))
            for r in h.runs:
                r.font.bold = True
        elif raw.startswith('## '):
            h_text = raw.replace('## ', '').strip()
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(4)
            add_formatted_runs(h, h_text, default_color=COLOR_SECONDARY, default_size=Pt(13))
            for r in h.runs:
                r.font.bold = True
        elif raw.startswith('### '):
            h_text = raw.replace('### ', '').strip()
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(3)
            add_formatted_runs(h, h_text, default_color=COLOR_PRIMARY, default_size=Pt(11))
            for r in h.runs:
                r.font.bold = True
        elif raw.startswith('> '):
            q_text = raw.replace('> ', '').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, q_text, default_color=COLOR_SECONDARY, default_size=Pt(9.5))
            for r in p.runs:
                r.font.italic = True
        elif raw.startswith('- ') or raw.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_runs(p, raw, default_color=COLOR_DARK, default_size=Pt(9.5), is_bullet=True)
        elif re.match(r'^\d+\.\s+', raw):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            item_text = re.sub(r'^\d+\.\s+', '', raw)
            add_formatted_runs(p, item_text, default_color=COLOR_DARK, default_size=Pt(9.5))
        elif raw in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run("─" * 45)
            r.font.color.rgb = COLOR_MUTED
            r.font.size = Pt(8)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            add_formatted_runs(p, raw, default_color=COLOR_DARK, default_size=Pt(9.5))

    if in_table:
        flush_table()

def generate_report(input_md, output_docx, author_name="Creador / Author"):
    if not os.path.exists(input_md):
        print(f"❌ Error: Input file not found: {input_md}")
        return

    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Portada Editorial Universal y Neutra
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(30)

    p_badge = doc.add_paragraph()
    r_badge = p_badge.add_run("💎 ARQUEÓLOGO CREATIVO / CREATIVE ARCHAEOLOGIST")
    r_badge.font.name = "Arial"
    r_badge.font.size = Pt(11)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_SECONDARY

    p_tit = doc.add_paragraph()
    r_tit = p_tit.add_run("ARQUEÓLOGO CREATIVO: RESCATA LOS DIAMANTES DE TU CAJÓN")
    r_tit.font.name = "Arial"
    r_tit.font.size = Pt(21)
    r_tit.font.bold = True
    r_tit.font.color.rgb = COLOR_PRIMARY
    p_tit.paragraph_format.space_after = Pt(8)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Auditoría Forense de Portafolio, Radiografía DAFO, Rescate de IPs y Roadmap de Monetización")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_MUTED
    p_sub.paragraph_format.space_after = Pt(35)

    # Tabla Metadatos Universal
    t_meta = doc.add_table(rows=3, cols=2)
    t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_info = [
        ("AUTOR / ESTUDIO AUDITADO:", author_name),
        ("TRIBUNAL AUDITOR:", "Panel de 6 Agentes (IP Scout, Crítico, Buyer Persona, Hater, Pricing, Operativo)"),
        ("METODOLOGÍA:", "Evidencia Longitudinal + Embudo de Activos + Grafo Circular + Matriz IE")
    ]
    for idx, (lbl, val) in enumerate(meta_info):
        c1, c2 = t_meta.cell(idx, 0), t_meta.cell(idx, 1)
        set_cell_margins(c1, 80, 80, 120, 120)
        set_cell_margins(c2, 80, 80, 120, 120)
        set_cell_background(c1, "EDF2F7")
        set_cell_background(c2, "F7FAFC")
        c1.paragraphs[0].paragraph_format.space_after = Pt(0)
        c2.paragraphs[0].paragraph_format.space_after = Pt(0)
        
        r1 = c1.paragraphs[0].add_run(lbl)
        r1.font.name = "Arial"
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = COLOR_PRIMARY
        
        r2 = c2.paragraphs[0].add_run(val)
        r2.font.name = "Arial"
        r2.font.size = Pt(9)
        r2.font.color.rgb = COLOR_DARK

    doc.add_page_break()

    parse_markdown_to_docx(doc, input_md)

    doc.save(output_docx)
    print(f"🎉 Informe Word maquetado generado con éxito: {output_docx}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Compilador universal de informes Word para Arqueólogo Creativo")
    parser.add_argument("--input", required=True, help="Ruta del archivo Markdown de entrada")
    parser.add_argument("--output", default="Informe_Arqueologo_Creativo.docx", help="Ruta del archivo Word de salida")
    parser.add_argument("--author", default="Creador Auditado", help="Nombre del autor o estudio")
    args = parser.parse_args()

    generate_report(args.input, args.output, args.author)
