import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def build_alby_master_docx():
    doc = Document()
    
    # Configuración de márgenes (1 pulgada)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    COLOR_PRIMARY = RGBColor(26, 36, 56)      # Navy Ejecutivo Profundo
    COLOR_SECONDARY = RGBColor(180, 115, 30)  # Oro / Bronce Arqueológico
    COLOR_DARK = RGBColor(45, 55, 72)
    COLOR_MUTED = RGBColor(113, 128, 150)
    COLOR_CODE_BG = "F7FAFC"

    # =========================================================================
    # 1. PORTADA EDITORIAL PREMIUM
    # =========================================================================
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(40)

    p_badge = doc.add_paragraph()
    r_badge = p_badge.add_run("💎 ARQUEÓLOGO CREATIVO / CREATIVE ARCHAEOLOGIST")
    r_badge.font.name = "Arial"
    r_badge.font.size = Pt(11)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_SECONDARY

    p_title = doc.add_paragraph()
    r_title = p_title.add_run("INFORME MAESTRO DE AUDITORÍA FORENSE & RESCATE DE ACTIVOS")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY
    p_title.paragraph_format.space_after = Pt(8)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Dictamen Estratégico, Desentierro de IPs, Radiografía Financiera y Roadmap 2026–2027")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_MUTED
    p_sub.paragraph_format.space_after = Pt(40)

    # Cuadro de metadata de la portada
    table_meta = doc.add_table(rows=4, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("AUTOR AUDITADO:", "Alby Ojeda (Guionista, Diseñador Narrativo & Creador)"),
        ("TRIBUNAL AUDITOR:", "Panel de 6 Agentes (IP Scout, Crítico, Buyer Persona, Hater, Pricing, Psicólogo)"),
        ("FECHA DE EMISIÓN:", "21 de Agosto de 2026"),
        ("ALCANCE DEL ANÁLISIS:", "+300.000 archivos | 522 guiones | 10 años de facturas históricas")
    ]
    for idx, (label, val) in enumerate(meta_data):
        c1 = table_meta.cell(idx, 0)
        c2 = table_meta.cell(idx, 1)
        c1.text = label
        c2.text = val
        set_cell_margins(c1, top=100, bottom=100, left=140, right=140)
        set_cell_margins(c2, top=100, bottom=100, left=140, right=140)
        set_cell_background(c1, "EDF2F7")
        set_cell_background(c2, "F7FAFC")
        c1.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        c1.paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY
        c2.paragraphs[0].runs[0].font.size = Pt(9.5)
        c2.paragraphs[0].runs[0].font.color.rgb = COLOR_DARK

    doc.add_page_break()

    # =========================================================================
    # 2. CARGA DE CONTENIDOS DESDE LOS INFORMES ESTRATÉGICOS
    # =========================================================================
    source_files = [
        r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\Estrategia\00_INFORME_MAESTRO_DEFINITIVO_Y_CALIBRADO.md",
        r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\Estrategia\02_LISTA_DE_PRECIOS_REALISTA_2026.md",
        r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\Estrategia\04_GUIA_MAESTRA_FILOSOFIA_LABORAL_Y_FILTRADO.md",
        r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\05_DICTAMEN_DEMOLEDOR_CONSEJO_MAESTRO.md",
        r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\2 Nivel Tactica\06_ROADMAP_ESTRATEGICO_PRODUCTOS_2026_2027.md"
    ]

    for s_file in source_files:
        if not os.path.exists(s_file):
            continue

        with open(s_file, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()

        in_table = False
        table_rows = []
        in_code_block = False

        for line in lines:
            raw = line.strip()
            if not raw:
                continue

            # Bloques de texto/cajas ASCII ```text ... ```
            if raw.startswith('```'):
                in_code_block = not in_code_block
                continue
            
            if in_code_block:
                p_code = doc.add_paragraph()
                p_code.paragraph_format.left_indent = Inches(0.3)
                p_code.paragraph_format.space_after = Pt(2)
                r_code = p_code.add_run(raw)
                r_code.font.name = "Consolas"
                r_code.font.size = Pt(8.5)
                r_code.font.color.rgb = COLOR_DARK
                continue

            # Tablas Markdown
            if raw.startswith('|') and raw.endswith('|'):
                if '---' in raw:
                    continue
                cells = [c.strip() for c in raw.split('|')[1:-1]]
                table_rows.append(cells)
                in_table = True
                continue
            elif in_table:
                if table_rows:
                    col_count = len(table_rows[0])
                    t = doc.add_table(rows=len(table_rows), cols=col_count)
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    t.autofit = True
                    for r_i, row in enumerate(table_rows):
                        for c_i, val in enumerate(row):
                            if c_i < col_count:
                                cell = t.cell(r_i, c_i)
                                cell.text = val
                                set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
                                if r_i == 0:
                                    set_cell_background(cell, "1A2438")
                                    for p in cell.paragraphs:
                                        for r in p.runs:
                                            r.font.name = "Arial"
                                            r.font.size = Pt(9)
                                            r.font.bold = True
                                            r.font.color.rgb = RGBColor(255, 255, 255)
                                else:
                                    if r_i % 2 == 1:
                                        set_cell_background(cell, "F8F9FA")
                                    for p in cell.paragraphs:
                                        for r in p.runs:
                                            r.font.name = "Arial"
                                            r.font.size = Pt(8.5)
                                            r.font.color.rgb = COLOR_DARK
                    doc.add_paragraph()
                table_rows = []
                in_table = False

            # Encabezados
            if raw.startswith('# '):
                h = doc.add_heading(raw.replace('# ', '').strip(), level=1)
                h.paragraph_format.space_before = Pt(22)
                h.paragraph_format.space_after = Pt(8)
                for r in h.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(16)
                    r.font.bold = True
                    r.font.color.rgb = COLOR_PRIMARY
            elif raw.startswith('## '):
                h = doc.add_heading(raw.replace('## ', '').strip(), level=2)
                h.paragraph_format.space_before = Pt(18)
                h.paragraph_format.space_after = Pt(6)
                for r in h.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(13)
                    r.font.bold = True
                    r.font.color.rgb = COLOR_SECONDARY
            elif raw.startswith('### '):
                h = doc.add_heading(raw.replace('### ', '').strip(), level=3)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(4)
                for r in h.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(11)
                    r.font.bold = True
                    r.font.color.rgb = COLOR_PRIMARY
            elif raw.startswith('> '):
                p_quote = doc.add_paragraph()
                p_quote.paragraph_format.left_indent = Inches(0.4)
                p_quote.paragraph_format.space_before = Pt(6)
                p_quote.paragraph_format.space_after = Pt(6)
                r_quote = p_quote.add_run(raw.replace('> ', '').strip())
                r_quote.font.name = "Arial"
                r_quote.font.size = Pt(10)
                r_quote.font.italic = True
                r_quote.font.color.rgb = COLOR_SECONDARY
            elif raw.startswith('- ') or raw.startswith('* '):
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(3)
                r_bullet = p_bullet.add_run(raw[2:].strip())
                r_bullet.font.name = "Arial"
                r_bullet.font.size = Pt(9.5)
                r_bullet.font.color.rgb = COLOR_DARK
            else:
                p_body = doc.add_paragraph()
                p_body.paragraph_format.space_after = Pt(6)
                r_body = p_body.add_run(raw)
                r_body.font.name = "Arial"
                r_body.font.size = Pt(9.5)
                r_body.font.color.rgb = COLOR_DARK

        doc.add_page_break()

    out_file = r"D:\PROYECTOS\_PANEL DE CONTROL_\_MASTER PLAN OPERATIVO_\INFORME_MAESTRO_ARQUEOLOGO_CREATIVO_ALBY_OJEDA.docx"
    doc.save(out_file)
    print(f"🎉 DOCUMENTO WORD MAQUETADO GENERADO CON ÉXITO: {out_file}")

if __name__ == "__main__":
    build_alby_master_docx()
